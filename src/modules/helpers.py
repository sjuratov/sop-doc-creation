import azure.cognitiveservices.speech as speechsdk
import base64
import cv2
import datetime
import json
import librosa
import matplotlib.pyplot as plt
import openai
import os
import pandas as pd
import requests
import sys
import time

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from IPython.display import Audio, Video, FileLink, Image
from mimetypes import guess_type
from moviepy.editor import VideoFileClip
from openai import AzureOpenAI
from pydub import AudioSegment
from wordcloud import WordCloud

from pytubefix import YouTube
from pytubefix.cli import on_progress

import streamlit as st

from dotenv import find_dotenv, load_dotenv

# Checking if the azd config file exists.
# If so, use it to source env variables
config_path = '../../.azure/config.json'
default_environment = None

if os.path.exists(config_path):
    with open(config_path, 'r') as config_file:
        config_data = json.load(config_file)
        default_environment = config_data.get('defaultEnvironment')
        if default_environment:
            print(f"Default Environment used: {default_environment}")
        else:
            print("defaultEnvironment parameter not found in the config file.")
else:
    print(f"Config file {config_path} does not exist. Not local execuriton or 'azd up' has not been executed.")

# If azd config file not found, use the standard .env file
if default_environment:
    load_dotenv(f"../../.azure/{default_environment}/.env",override=True)
else:
    load_dotenv(find_dotenv(),override=True)

# Azure Speech services
AZURE_SPEECH_KEY = os.getenv("AZURE_SPEECH_KEY")
AZURE_SPEECH_REGION = os.getenv("AZURE_SPEECH_REGION")

# Azure OpenAI
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")

print(f"Azure OpenAI endpoint (helpers): {AZURE_OPENAI_ENDPOINT}")

def download_file(url, path):
    #import requests
    #import os

    if not os.path.exists(os.path.dirname(path)):
        try:
            os.makedirs(os.path.dirname(path))
        except Exception as e:
            print(f"Error: {e}")

    try:
        response = requests.get(url, stream=True)
        with open(path, "wb") as file:
            for chunk in response.iter_content(chunk_size=128):
                file.write(chunk)
    except Exception as e:
        print(f"Error: {e}")

def download_youtube_video(url, path):
    # from pytubefix import YouTube
    # from pytubefix.cli import on_progress
    # import os

    if not os.path.exists(os.path.dirname(path)):
        try:
            os.makedirs(os.path.dirname(path))
        except Exception as e:
            print(f"Error: {e}")

    try:
        youtube = YouTube(url, on_progress_callback = on_progress)
        video = youtube.streams.get_highest_resolution()
        video.download(path)
        return video
    except Exception as e:
        print(f"Error: {e}")

def display_file_info(file_name):
    """
    Display the size, and last modification date and time of a specified file.

    Parameters:
    file_path (str): The path to the file whose information is to be displayed.

    Returns:
    file_name, file_size, formatted_time
    """
    
    if not os.path.isfile(file_name):
        print(f"The file at {file_name} does not exist.")
        return

    # Get the file size
    file_size_bytes = os.path.getsize(file_name)
    file_size_mb = file_size_bytes / (1024 * 1024)
    
    # Get the last modification time
    modification_time = os.path.getmtime(file_name)
    
    # Convert modification time to a readable format
    formatted_time = datetime.datetime.fromtimestamp(modification_time).strftime('%Y-%m-%d %H:%M:%S')
    
    # Display the file information
    print(f"File: {file_name}")
    print(f"Size: {file_size_mb:.2f} MB")
    print(f"Last Modified: {formatted_time}")
    
    return file_name, file_size_mb, formatted_time

def display_video_frames(video_file):
    """
    Display a specified number of equally time-distributed frames from a video file in their temporal order.

    Parameters:
    video_path (str): Path to the MP4 video file.
    num_frames (int): Number of equally distributed frames to display. Default is 10.

    Returns:
    None
    """
    
    num_frames=10
    
    # Open the video file
    cap = cv2.VideoCapture(video_file)
    
    # Check if the video was opened successfully
    if not cap.isOpened():
        print("Error: Could not open video file.")
        return

    # Get the total number of frames
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    
    if total_frames < num_frames:
        print(f"Error: The video contains only {total_frames} frames, which is less than {num_frames}.")
        cap.release()
        return

    # Calculate interval to select frames
    interval = total_frames // num_frames
    
    # Generate equally distributed frame indices
    frame_indices = [i * interval for i in range(num_frames)]

    # Set up the plot for displaying frames
    plt.figure(figsize=(15, 10))
    
    for i, frame_idx in enumerate(frame_indices):
        # Set the video position to the frame index
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        
        # Read the frame
        ret, frame = cap.read()
        
        if not ret:
            print(f"Error: Could not read frame at index {frame_idx}.")
            continue
        
        # Convert frame from BGR to RGB
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        
        # Plot the frame
        plt.subplot(2, 5, i + 1)  # Adjust subplot grid as needed
        plt.imshow(frame_rgb)
        plt.title(f"Frame {frame_idx}")
        plt.axis('off')
    
    # Show all frames
    plt.tight_layout()
    #plt.show()
    # st.divider()
    st.markdown("#### Video Frames")
    st.pyplot(plt)
    
    # Release the video capture object
    cap.release()

def get_audio_file(video_file, RESULTS_DIR):
    """   
    Extracts the audio track from a given video file and saves it as a WAV file.

    This function takes the path to a video file, extracts the audio track using the
    `moviepy` library, and saves the extracted audio as a `.wav` file in the specified
    audio directory. It prints the start and completion of the extraction process and
    reports the elapsed time taken for the operation.

    Parameters:
    video_file (str): The full path to the input video file from which audio will be extracted.
                      The file should be in a format supported by `moviepy`, such as MP4 or MKV.

    Returns:
    str: The full path to the saved audio file in WAV format. The path includes the directory
         specified by `AUDIO_DIR` and the filename derived from the input video file.
    """
    
    print("Audio extraction from video file {video_file}\n")
    start = time.time()

    audio_file = os.path.join(
        RESULTS_DIR,
        os.path.splitext(os.path.basename(video_file))[0] + ".wav")
    
    # Loading video file
    video_clip = VideoFileClip(video_file)
    audio_clip = video_clip.audio
    
    # Saving audio file
    audio_clip.write_audiofile(audio_file)
    audio_clip.close()
    video_clip.close()
    
    # End
    print("\nDone")
    elapsed = time.time() - start
    print("Elapsed time: " + time.strftime(
        "%H:%M:%S.{}".format(str(elapsed % 1)[2:])[:15], time.gmtime(elapsed)))

    return audio_file

def display_amplitude(audio_file):
    """
    This function reads an audio file, computes its amplitude waveform, and visualizes
    it using a waveplot. The plot shows the variation in amplitude over time, providing
    a graphical representation of the audio signal.

    Parameters:
    audio_file (str): The full path to the audio file to be visualized. The file should be
                      in a format supported by `librosa`, such as WAV or MP3.

    Returns:
    None: This function does not return any value. It displays a matplotlib figure showing
          the amplitude of the audio signal.
    """
    
    plt.figure(figsize=(15, 5))
    
    # Reading the sound file
    y, sr = librosa.load(audio_file)
    
    # Amplitude plot
    librosa.display.waveshow(y, sr=sr)
    title = f"Waveplot of {audio_file}"
    plt.title(title, fontdict=dict(size=15))
    plt.xlabel("Time", fontdict=dict(size=12))
    plt.ylabel("Amplitude", fontdict=dict(size=12))

    # plt.show()
    st.pyplot(plt)

def azure_text_to_speech(audio_filepath, locale, disp=False):
    """
    Transcribes speech from an audio file using Azure Speech-to-Text (TTS) service.

    This function sends an audio file to the Azure Speech-to-Text service for transcription. 
    It configures the Azure speech recognizer with the specified locale and processes the 
    audio file to extract the transcription text, confidence scores, and word-level details. 
    The function also measures and prints the time taken for the transcription process.

    Parameters:
    audio_filepath (str): The full path to the audio file to be transcribed. The audio file 
                          should be in a format supported by Azure Speech-to-Text service.
    locale (str): The language and region code for the transcription, e.g., 'en-US' for 
                  English (United States). This specifies the language model to be used for 
                  transcription.
    disp (bool, optional): If set to True, the function will print the transcription results, 
                            confidence scores, and word-level details. Defaults to False.

    Returns:
    tuple: A tuple containing three lists:
        - transcript_display_list (list): List of transcriptions as displayed text.
        - confidence_list (list): List of confidence scores corresponding to the transcriptions.
        - words (list): List of words with their details, including timing and confidence.
    """

    print(f"Running Speech to text from audio file {audio_filepath}\n")
    start = time.time()

    # Config
    audio_config = speechsdk.audio.AudioConfig(filename=audio_filepath)
    speech_config = speechsdk.SpeechConfig(subscription=AZURE_SPEECH_KEY,
                                           region=AZURE_SPEECH_REGION)
    # Timestamps are required
    speech_config.request_word_level_timestamps()
    speech_config.speech_recognition_language = locale
    speech_config.output_format = speechsdk.OutputFormat(1)

    # Creates a recognizer with the given settings
    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config,
                                                   audio_config=audio_config)

    # Variable to monitor status
    done = False

    # Service callback for recognition text
    transcript_display_list = []
    transcript_ITN_list = []
    confidence_list = []
    words = []

    def parse_azure_result(evt):
        import json

        response = json.loads(evt.result.json)
        transcript_display_list.append(response["DisplayText"])
        confidence_list_temp = [
            item.get("Confidence") for item in response["NBest"]
        ]
        max_confidence_index = confidence_list_temp.index(
            max(confidence_list_temp))
        confidence_list.append(
            response["NBest"][max_confidence_index]["Confidence"])
        transcript_ITN_list.append(
            response["NBest"][max_confidence_index]["ITN"])
        words.extend(response["NBest"][max_confidence_index]["Words"])

    # Service callback that stops continuous recognition upon receiving an event `evt`
    def stop_cb(evt):
        print("CLOSING on {}".format(evt))
        speech_recognizer.stop_continuous_recognition()
        nonlocal done
        done = True

        if disp:
            # Do something with the combined responses
            print(transcript_display_list)
            print(confidence_list)
            print(words)

    # Connect callbacks to the events fired by the speech recognizer
    speech_recognizer.recognizing.connect(
        lambda evt: logger.debug("RECOGNIZING: {}".format(evt)))
    speech_recognizer.recognized.connect(parse_azure_result)
    speech_recognizer.session_started.connect(
        lambda evt: logger.debug("SESSION STARTED: {}".format(evt)))
    speech_recognizer.session_stopped.connect(
        lambda evt: logger.debug("SESSION STOPPED {}".format(evt)))
    speech_recognizer.canceled.connect(
        lambda evt: logger.debug("CANCELED {}".format(evt)))
    # stop continuous recognition on either session stopped or canceled events
    speech_recognizer.session_stopped.connect(stop_cb)
    speech_recognizer.canceled.connect(stop_cb)

    # Start continuous speech recognition
    speech_recognizer.start_continuous_recognition()
    while not done:
        time.sleep(0.5)

    print("\nDone")
    elapsed = time.time() - start
    print("Elapsed time: " + time.strftime(
        "%H:%M:%S.{}".format(str(elapsed % 1)[2:])[:15], time.gmtime(elapsed)))

    return transcript_display_list, confidence_list, words

def text_search(text):
    """
    Searches for a specific word in a DataFrame and retrieves its offset and duration.

    This function searches for a given word in a DataFrame containing transcription data. 
    It returns the rows of the DataFrame that match the specified word, along with the 
    offset and duration associated with that word. The offset and duration are extracted 
    from the DataFrame for the first occurrence of the word.

    Parameters:
    text (str): The word or phrase to search for in the DataFrame. The search is case-sensitive 
                and matches exactly.

    Returns:
    tuple: A tuple containing:
        - df_res (DataFrame): The rows of the DataFrame where the 'Word' column matches the 
                              specified text.
        - offset_secs (float): The offset time in seconds for the first occurrence of the word.
        - duration_secs (float): The duration in seconds for the first occurrence of the word.

    """

    # Search
    df_res = df[df["Word"] == text]
    offset_secs = df_res["Offset_in_secs"].values[0]
    duration_secs = df_res["Duration_in_secs"].values[0]

    return df_res, offset_secs, duration_secs

def ask_gpt4o(prompt, sop_text, model):
    """
    Sends a prompt to the GPT-4 model via Azure OpenAI and returns the response.

    This function uses the Azure OpenAI service to send a specified prompt to the GPT-4o model.
    The function is designed to analyze an audio transcript file and return the output in a JSON file format.

    Args:
        prompt (str): The input prompt to be sent to the GPT-4o model.

    Returns:
        str: The content of the response from the GPT-4o model.
    """
    
    # Azure OpenAI client
    client = AzureOpenAI(azure_endpoint=AZURE_OPENAI_ENDPOINT,
                         api_key=AZURE_OPENAI_KEY,
                         api_version="2024-02-01")
    
    # client = AzureOpenAI(azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    #                      api_key=os.getenv("AZURE_OPENAI_KEY"),
    #                      api_version="2024-02-01")

    # Response with the json object property
    response = client.chat.completions.create(
        model=model,
        response_format={"type": "json_object"},
        temperature=0.0,
        max_tokens=2000,
        messages=[
            {
                "role":
                "system",
                "content":
                "You are an AI assistant that is analysing an audio transcript file.\
             Print the output in a JSON file format."
            },
            {
                "role": "user",
                "content": f"{prompt}: {sop_text}"
            },
        ],
    )

    return response.choices[0].message.content

def get_video_info(video_file):
    """
    Prints the length, number of frames, and frames per second (FPS) of a video file.

    This function opens a video file using OpenCV, retrieves information about the total number of frames,
    frames per second (FPS), and calculates the duration of the video in hours, minutes, and seconds. It then
    prints this information.

    Parameters:
    video_file (str): Path to the video file.

    Returns:
    duration, total of frames and fps
    """
    
    # Open the video file
    cap = cv2.VideoCapture(video_file)
    print(f"Video file: {video_file}")
    
    if not cap.isOpened():
        print(f"Error: Cannot open video file {video_file}")
        return

    # Get the total number of frames
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    # Get the frames per second (FPS)
    fps = cap.get(cv2.CAP_PROP_FPS)
    # Calculate the duration of the video in seconds
    duration = total_frames / fps
    # Convert duration to hours, minutes, and seconds
    hours = int(duration // 3600)
    minutes = int((duration % 3600) // 60)
    seconds = int(duration % 60)

    # Print the video information
    print(f"\n- Duration in seconds: {duration:.0f}")
    print(f"- Length of video: {hours:02}:{minutes:02}:{seconds:02}")
    print(f"- Number of frames: {total_frames}")
    print(f"- Frames per second (FPS): {fps:.0f}")

    # Release the video capture object
    cap.release()
    
    return duration, total_frames, fps

def get_video_frame(video_file, offset_in_secs, FRAMES_DIR):
    """
    Extracts a frame from a video file at a specified offset in seconds and saves it as an image file.

    Args:
        video_file (str): Path to the video file.
        offset_in_secs (float): The offset in seconds from which to capture the frame.

    Returns:
        str: Path to the saved frame image file.
    """
    
    # Open the video file
    cap = cv2.VideoCapture(video_file)

    # Check if the video opened successfully
    if not cap.isOpened():
        print("Error: Could not open video.")
        return

    # Get the frames per second (fps) of the video
    fps = cap.get(cv2.CAP_PROP_FPS)

    # Calculate the frame number to capture
    frame_number = int(offset_in_secs * fps)

    # Set the video to start at the calculated frame
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)

    # Read the frame
    ret, frame = cap.read()

    if not ret:
        print("Error: Could not read frame.")
        return

    # Save the frame as an image
    frame_file = os.path.join(
        FRAMES_DIR,
        f"{os.path.splitext(os.path.basename(video_file))[0]}_frame_{str(offset_in_secs)}.png"
    )
    cv2.imwrite(frame_file, frame)

    # Release the video capture object
    cap.release()

    return frame_file

def local_image_to_data_url(image_path):
    """
    Convert a local image file to a data URL.

    This function takes the path of a local image file, reads the file,
    encodes its contents in base64, and constructs a data URL with the
    appropriate MIME type.

    Args:
        image_path (str): The file path to the local image.

    Returns:
        str: A data URL containing the base64-encoded image data.

    """
    
    mime_type, _ = guess_type(image_path)

    if mime_type is None:
        mime_type = "application/octet-stream"

    with open(image_path, "rb") as image_file:
        base64_encoded_data = base64.b64encode(image_file.read()).decode("utf-8")

    return f"data:{mime_type};base64,{base64_encoded_data}"

def gpt4o_imagefile(image_file, prompt, model):
    """
    Analyze an image file using Azure OpenAI's GPT-4 model.

    This function sends an image and a text prompt to Azure OpenAI's GPT-4 model
    for analysis. The image is first converted to a data URL and then included
    in the message payload along with the text prompt. The function returns the
    response from the GPT-4 model.

    Args:
        image_file (str): The file path to the local image.
        prompt (str): The text prompt to accompany the image for analysis.

    Returns:
        dict: The response from Azure OpenAI's GPT-4 model containing the analysis results.
    """
    
    client = AzureOpenAI(
        api_key=AZURE_OPENAI_KEY,
        api_version="2024-05-01-preview",
        base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{model}",
    )

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": "You are an AI helpful assistant to analyse images.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": local_image_to_data_url(image_file)},
                    },
                ],
            },
        ],
        max_tokens=2000,
        temperature=0.0,
    )

    return response

def checklist_docx_file(video_file, json_data, RESULTS_DIR, model, nb_images_per_step=3):
    """
    Generates a DOCX file containing a checklist based on video frames and provided JSON data.

    This function creates a Word document with a checklist where each checklist step includes
    a heading, summary, keywords, and images extracted from a video file at specified offsets.

    Args:
        video_file (str): Path to the video file from which frames are extracted.
        json_data (list of dict): List of dictionaries where each dictionary represents a checklist step
                                  containing keys like 'Step', 'Summary', 'Keywords', 'Offset', and 'Offset_in_secs'.
        nb_images_per_step (int, optional): Number of images to include for each checklist step. Defaults to 3.

    Returns:
        str: Path to the generated DOCX file.
    """
    
    print("Generating checklist file...")
    
    image_size = 5 # size of each image that will be inserted

    FRAMES_DIR = f"{RESULTS_DIR}/frames"
    
    # Filename
    docx_file = os.path.join(
        RESULTS_DIR,
        os.path.splitext(os.path.basename(video_file))[0] + ".docx")

    # Initialize the document
    doc = Document()

    # Adding a header for each page
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = f"SOP document"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Heading level 1
    doc.add_heading(f"Checklist document for video: {video_file}", level=1)
    doc.add_paragraph("")
    
    duration = 0  # do not change
    
    # Process each step from the JSON data
    for idx, step in enumerate(json_data, start=1):
        # get values
        title = str(step['Title']).upper()
        summary = step['Summary']
        keywords = step['Keywords']
        offset_secs = step['Offset_in_secs']
        duration = round(step['Offset_in_secs'] - duration, 3)
        
        # Add checklist step details to the document
        doc.add_heading(f"{idx} Checklist step {step['Step']}: {title}", level=2)
        doc.add_paragraph("")
        doc.add_paragraph(f"Summary: {summary}")
        doc.add_paragraph(f"Keywords: {keywords}")
        doc.add_paragraph(f"Offset in seconds: {offset_secs}")
        doc.add_paragraph(f"Duration in seconds: {duration}")

        # Add images & automatic caption for the current step
        for img_idx in range(1, nb_images_per_step + 1):
            # Retrieve the frame and add it to the document
            frame_file = get_video_frame(
                video_file,
                int(offset_secs) + img_idx * 3,
                FRAMES_DIR)            
            doc.add_picture(frame_file, width=Inches(image_size))
        
            # Adding the automatic caption of the frame
            caption_image = gpt4o_imagefile(frame_file, "Generate a detailled caption of this image.", model)
            caption = caption_image.choices[0].message.content
            doc.add_paragraph(f"- Automatic frame caption: {caption}")
            
            # OCR of the frame
            ocr_image = gpt4o_imagefile(frame_file, "Print all the extracted text from this image separated with a comma", model)
            ocr = ocr_image.choices[0].message.content
            doc.add_paragraph(f"- Automatic OCR: {ocr}")
            
            # Deleting the frame file (optional)
            os.remove(frame_file)
    
        # Add a blank line for spacing
        doc.add_page_break()

    # Adding a footnote
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    now = str(datetime.datetime.today().strftime('%d-%b-%Y'))
    footer_para.text = f"{now} | Powered by Azure AI services"

    # Save the document
    doc.save(docx_file)
    
    # End
    print(f"\nDone. Checklist file has been saved to {docx_file}")

    return docx_file

